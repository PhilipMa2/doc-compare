import textract
import docx 
import re 
import difflib
import sys
from pypdf import PdfReader

# Find the first entry by skipping table of contents
def find_2nd(string, substring, start):
    return string.find(substring, start + 1)

# Return list of strings such that each string is one entry in the document
def preprocess(name, sections):
    format = name[name.rfind('.')+1:]
    text = ''
    if format == 'docx':
        file = textract.process(name)
        text = file.decode()
    elif format == 'pdf':
        file = PdfReader(name)
        for page in file.pages:
            content = page.extract_text()
            match = re.search('\D', content)
            text += ' ' + content[match.start():]
    else:
        raise ValueError('The files needs to be a docx or pdf.')

    start = text.find("第一章")
    if start == -1:
        start = text.find("第一条")
        if start == -1:
            raise ValueError('The format of file given is not supported.')
    second_occurence = text.find("第一章", start + 1)
    if second_occurence != -1:
        text = text[second_occurence:]
    else:
        text = text[start:]
    texts = re.findall(r'第.{1,5}[章|条][\s\S]*?\s(?=第.+[章|条])', text)
    last_match = re.search(r'条.{1,5}第', text[::-1])
    texts.append(text[len(text) - last_match.end():])
    for i in range(len(texts)):
        if re.search('^第.{1,5}章', texts[i]) is not None:
            sections.append(texts[i])
    return texts

def only_add_new(oldp, newp):
    run = newp.add_run(new_texts[j])
    run.font.color.rgb = docx.shared.RGBColor(128,21,0)
    run = oldp.add_run('新 增')
    run.font.color.rgb = docx.shared.RGBColor(0,102,51)

def only_add_old(oldp, newp):
    run = oldp.add_run(old_texts[i])
    run.font.color.rgb = docx.shared.RGBColor(0,102,51)
    run = newp.add_run('移 除')
    run.font.color.rgb = docx.shared.RGBColor(128,21,0)

def different(old_text, new_text):
    seq_mat = difflib.SequenceMatcher()
    seq_mat.set_seqs(old_text, new_text)
    ratio = seq_mat.ratio()
    len_diff = abs(len(old_text) - len(new_text))
    # if the lengths of two entries are substantially different, then it's 
    # possible for them to be the same entry even if difference ratio is less
    # than 0.5. For example, see entry 113 (2021) and 108 (2014)
    return (ratio < 0.5 and len_diff < 200) or (len_diff >= 200 and ratio < 0.1)

def different_section(old_sec, new_sec):
    seq_mat = difflib.SequenceMatcher()
    seq_mat.set_seqs(old_sec, new_sec)
    ratio = seq_mat.ratio()
    # len_diff = abs(len(old_sec) - len(new_sec))

    return ratio < 0.7

def add_both(oldp, newp, old_text, new_text, old_text_next=None, new_text_next=None):
    diff = difflib.ndiff(old_text, new_text)
    # find the difference ratio of two entries
    if not basic and (different(old_text, new_text) and (not old_text_next or not new_text_next or different(old_text_next, new_text_next))):
            only_add_new(oldp, newp)
            return False
    else:
        for k, s in enumerate(diff):
            # same content
            if s[0] == ' ': 
                newp.add_run(s[-1])
                oldp.add_run(s[-1])
            # if content is added in 2021, make it red in 2021
            elif s[0] == '+':
                run = newp.add_run(s[-1])
                run.font.color.rgb = docx.shared.RGBColor(128,21,0)
            # if content is deleted in 2021, make it green in 2014
            elif s[0] == '-':
                run = oldp.add_run(s[-1])
                run.font.color.rgb = docx.shared.RGBColor(0,102,51)
        return True
    
def match_section(keys, values):
    section_dict = {}
    prev_match = 0
    for key in keys:
        # print(key)
        for i in range(prev_match, len(values)):
            if not different_section(key, values[i]):
                section_dict[key] = values[i]
                prev_match = i + 1
                break
        if key not in section_dict:
            section_dict[key] = values[prev_match]
            prev_match += 1
        # print(section_dict[key])
        # print(different_section(key, section_dict[key]))
        # print('end')
    return section_dict


if __name__ == "__main__":
    basic = False
    if len(sys.argv) == 3:
        old_pos = 1
        new_pos = 2
    elif len(sys.argv) == 4:
        old_pos = 2
        new_pos = 3
        basic = True
    else:
        print("USAGE: python3 compare.py [-basic] <old_file.[docx|pdf]> <new_file.[docx|pdf]>")

        sys.exit(1)

    # if sys.argv[1][-4:] != 'docx' or sys.argv[2][-4:] != 'docx':
    #     print('Files need to be in docx format')
    #     sys.exit(1)

    # specify file name to compare
    try:
        old_sections = []
        new_sections = []
        old_texts = preprocess(sys.argv[old_pos], old_sections)
        new_texts = preprocess(sys.argv[new_pos], new_sections)
    except ValueError as err:
        print(err.args[0])
        sys.exit(1)

    if len(old_sections) <= len(new_sections):
        section_dict = match_section(old_sections, new_sections)
    else:
        section_dict = match_section(new_sections, old_sections)

    # create output document for comparison
    output = docx.Document()
    table = output.add_table(rows=1, cols=2, style="Table Grid")
    row = table.rows[0]
    # get the year of each document
    year_pos = old_texts[-1].find('年')
    old_year = old_texts[-1][year_pos-4:year_pos]
    year_pos = new_texts[-1].find('年')
    new_year = new_texts[-1][year_pos-4:year_pos]
    row.cells[0].text = old_year
    row.cells[1].text = new_year
    # i keeps track of old document law and j keeps track of new document law
    i = 0
    j = 0
    while i < len(old_texts) and j < len(new_texts):
        cells = table.add_row().cells
        oldp = cells[0].add_paragraph()
        newp = cells[1].add_paragraph()
        section_pattern = re.compile("^第.{1,5}章[\s\S]*")
        old_section_end = section_pattern.match(old_texts[i])
        new_section_end = section_pattern.match(new_texts[j])
        if old_section_end and not new_section_end:
            only_add_new(oldp, newp)
            j += 1
        elif not old_section_end and new_section_end:
            only_add_old(oldp, newp)
            i += 1    
        elif old_section_end and new_section_end:
            if len(old_sections) <= len(new_sections):
                if section_dict[old_texts[i]] == new_texts[j]:
                    add_both(oldp, newp, old_texts[i], new_texts[j])
                    i += 1
                    j += 1
                else:
                    only_add_new(oldp, newp)
                    j += 1
            elif section_dict[new_texts[j]] == old_texts[i]:
                add_both(oldp, newp, old_texts[i], new_texts[j])
                i += 1
                j += 1
            else:
                only_add_old(oldp, newp)
                i += 1
        else:
            if i < len(old_texts) - 1 and j < len(new_texts) - 1:
                both_added = add_both(oldp, newp, old_texts[i], new_texts[j], old_texts[i+1], new_texts[j+1])  
            else:
                both_added = add_both(oldp, newp, old_texts[i], new_texts[j]) 
            if both_added:
                i += 1
            j += 1
    while i < len(old_texts):
        cells = table.add_row().cells
        oldp = cells[0].add_paragraph()
        newp = cells[1].add_paragraph()
        only_add_old(oldp, newp)
        i += 1
    while j < len(new_texts):
        cells = table.add_row().cells
        oldp = cells[0].add_paragraph()
        newp = cells[1].add_paragraph()
        only_add_new(oldp, newp)
        j += 1
        
    output.save("output.docx")