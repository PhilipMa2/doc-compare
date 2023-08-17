from pypdf import PdfReader
import re
import docx
import difflib
import tabula
import pandas as pd

def increase_left(sec):
    pos = sec.find('\\')
    return str(int(sec[:pos])+1) 

def increase_right(sec):
    pos = sec.rfind('\\')
    return sec[:pos+2] + str(int(sec[pos+2:])+1)

def sync_mid(sec2, sec3):
    pos = sec3.find('\\')
    mid = sec2[sec2.find('.')+1:]
    return sec3[:pos+2] + mid

def preprocess(text):
    text = text[re.search('1 \S', text).start():]
    texts = []
    delimiter = ['2','1\.1','1\.1\.1']
    next_section = re.search('[^表\.\n](?:' + '|'.join(delimiter) + ') \S', text)

    while next_section:
        curr_section = text[:next_section.start()+1]
        texts.append(curr_section)
        text = text[next_section.start()+1:]
        section_num = text[:text.find(' ')]
        counts = section_num.count('.')
        if counts == 0:
            delimiter[0] = str(int(delimiter[0])+1)
            delimiter[1] = increase_left(delimiter[1]) + '\.1'
            delimiter[2] = increase_left(delimiter[2]) + '\.1\.1'
        elif counts == 1:
            delimiter[2] = sync_mid(delimiter[1], delimiter[2]) + '\.1'
            delimiter[1] = increase_right(delimiter[1])
        else:
            delimiter[2] = increase_right(delimiter[2])
        next_section = re.search('[^表\.\n](?:' + '|'.join(delimiter) + ') \S', text)
    return texts

def get_year(pdf):
    title = pdf.pages[0].extract_text()
    pos = title.find('GB')
    return title[pos:pos+11]

def compare_entry(oldp, newp, old_text, new_text):
    old_pos = old_text.find('见表')
    new_pos = new_text.find('见表')

    if old_pos != -1 and new_pos != -1:
        old_text = old_text[:old_pos+3]
        new_text = new_text[:new_pos+3]

    diff = difflib.ndiff(old_text, new_text)
    for k, s in enumerate(diff):
        # same content
        if s[0] == ' ': 
            newp.add_run(s[-1])
            oldp.add_run(s[-1])
        # if content is added in 2021, make it red in 2021
        elif s[0] == '+':
            run = newp.add_run(s[-1])
            run.font.color.rgb = docx.shared.RGBColor(128,21,0)
        # if content is deleted in 2021, make it green in 2014
        elif s[0] == '-':
            run = oldp.add_run(s[-1])
            run.font.color.rgb = docx.shared.RGBColor(0,102,51)

def get_table(file, page_num, page_text, all_dfs):
    dfs = tabula.read_pdf(file, pages=page_num)
    if len(dfs) > 0:
        i = 0
        if page_text.find('(续)') != -1:
            if len(dfs[0].columns) > len(all_dfs[-1].columns):
                dfs[0] = dfs[0].dropna(axis=1, thresh=2)
            new_cols = {x: y for x, y in zip(dfs[0].columns, all_dfs[-1].columns)}
            all_dfs[-1] = pd.concat([all_dfs[-1], dfs[i].rename(columns=new_cols)], ignore_index=True)
            i += 1
        all_dfs.extend(dfs[i:])


if __name__ == "__main__":
    old_file = 'A.pdf'
    new_file = 'B.pdf'
    old_pdf = PdfReader(old_file)
    new_pdf = PdfReader(new_file)
    old_text = ''
    new_text = ''
    old_year = get_year(old_pdf)
    new_year = get_year(new_pdf)
    old_dfs = []
    new_dfs = []
    
    for i in range(5):
        text = old_pdf.pages[i].extract_text()
        get_table(old_file, i+1, text, old_dfs)
        old_text += text
    for i in range(5):
        text = new_pdf.pages[i].extract_text()
        get_table(new_file, i+1, text, new_dfs)
        new_text += text

    # old_texts = preprocess(old_text)
    # new_texts = preprocess(new_text)

    output = docx.Document()
    # table = output.add_table(rows=1, cols=2, style="Table Grid")
    # row = table.rows[0]
    # # get the year of each document
    # row.cells[0].text = old_year
    # row.cells[1].text = new_year

    # # i keeps track of old document law and j keeps track of new document law
    # i = 0
    # j = 0
    # while i < len(old_texts) and j < len(new_texts):
    #     cells = table.add_row().cells
    #     oldp = cells[0].add_paragraph()
    #     newp = cells[1].add_paragraph()
    #     compare_entry(oldp, newp, old_texts[i], new_texts[j])
    #     i += 1
    #     j += 1
    df_all = old_dfs[0].join(new_dfs[0], how='outer', lsuffix='_old', rsuffix='_new')
    table = output.add_table(df_all.shape[0]+1, df_all.shape[1])
    for j in range(df_all.shape[-1]):
        table.cell(0, j).text = df_all.columns[j]
    i = j = 0
    wait_i = wait_j = False
    for k in range(df_all.shape[0]):
        if i < old_dfs[0].shape[0] and not wait_i:
            table.cell(k+1, 0).text = str(df_all.values[i, 0])
            table.cell(k+1, 1).text = str(df_all.values[i, 1])
            if pd.isna(df_all.values[i, 1]):
                print('wait')
                if not wait_j:
                    wait_i = True
                else:
                    wait_j = False
            i += 1
        else:
            table.cell(k+1, 0).text = 'nan'
            table.cell(k+1, 1).text = 'nan'
        if j < new_dfs[0].shape[0] and not wait_j:
            table.cell(k+1, 2).text = str(df_all.values[j, 2])
            table.cell(k+1, 3).text = str(df_all.values[j, 3])
            if pd.isna(df_all.values[j, 3]):
                if not wait_i:
                    wait_j = True
                else:
                    wait_i = False
            j += 1
        else:
            table.cell(k+1, 2).text = 'nan'
            table.cell(k+1, 3).text = 'nan'
        
    output.save("output.docx")

        
    # with open('out.txt', 'w') as f:
    #     for t in old_texts:
    #         f.write(t)
    #         f.write('\n')
    #         f.write('\n')

    #     for t in new_texts:
    #         f.write(t)
    #         f.write('\n')
    #         f.write('\n')
    